# -*- coding: utf-8 -*-
"""IR -> .docx, default theme only.

This is the D2 version of the renderer: no company template, no drawn diagrams. It exists so the
feature reaches a point where you can double-click a real Word file and read it. D4 extends this
same file to load a template and fill the "Kiem soat tai lieu" table.

Contract
    render.py <payload.json>     write the documents described by the payload
    render.py --check            report whether this interpreter can render at all

Both modes print ONE json object on stdout and nothing else, so the Node side can parse it.

The traps below were paid for once already (RULESET.md section 6) - do not "simplify" them away:
  * tables must be forced to fixed layout, otherwise one long cell collapses every other column;
  * List Number shares a counter across the whole document, so numbered steps are written by hand;
  * writing into a cell's runs[0] inherits Normal - write into a run that has content;
  * a file open in Word raises PermissionError: report the name, never abort the whole export.
"""
import json
import os
import sys

OUT = {"ok": False, "files": [], "skipped": [], "warnings": [], "counts": {}}


def die(message, kind="error"):
    OUT["ok"] = False
    OUT["error"] = message
    OUT["errorKind"] = kind
    sys.stdout.write(json.dumps(OUT, ensure_ascii=False))
    sys.exit(0)      # a handled failure is data, not a crash: exit 0 keeps the JSON contract


try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:                                    # noqa: BLE001
    die("Thiếu gói python-docx: %s" % exc, kind="missing-docx")


# ---------------------------------------------------------------- theme (RULESET section 3)
ACCENT = "1F4E79"
ZEBRA = "F2F5F9"
CODE_BG = "F3F3F3"
LINK = "0563C1"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"
HEADING_PT = {1: 17.0, 2: 13.5, 3: 12.0, 4: 11.5}


def rgb(hex6):
    return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))


def shade(element, hex6):
    """Paint a paragraph or a table cell. Works on any element that accepts w:shd."""
    pr = element.get_or_add_tcPr() if element.tag.endswith("}tc") else element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    pr.append(shd)


def left_border(paragraph, hex6):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex6)
    borders.append(left)
    pPr.append(borders)


def apply_theme(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11.5)
    # East Asian font has its own slot; without it Word substitutes for Vietnamese diacritics.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.25
    pf.space_after = Pt(7)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in HEADING_PT.items():
        st = doc.styles["Heading %d" % level]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(ACCENT)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = BODY_FONT
    title.font.size = Pt(26)
    title.font.color.rgb = rgb(ACCENT)

    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)


def printable_width(doc):
    """Inches available between the margins. Length arithmetic in python-docx yields a bare int
    (EMU), not a Length, so the conversion is done here once instead of at every call site."""
    s = doc.sections[0]
    return (s.page_width - s.left_margin - s.right_margin) / 914400.0


# ---------------------------------------------------------------- fields, links, header/footer
def add_field(paragraph, instruction, placeholder=""):
    """A Word field (TOC, PAGE, NUMPAGES). Word fills it in on open / on F9."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, text, end):
        run._r.append(el)
    return run


def add_hyperlink(paragraph, url, label):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(underline)
    run.append(rPr)
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = label
    run.append(text)
    link.append(run)
    paragraph._p.append(link)


def build_header_footer(doc, meta):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    head = section.header.paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head.paragraph_format.space_after = Pt(0)
    bits = [meta["title"]]
    if meta.get("draft"):
        bits.append("BẢN NHÁP")
    run = head.add_run("  ·  ".join(bits))
    run.font.size = Pt(8)
    run.font.name = BODY_FONT
    run.font.color.rgb = rgb("808080")

    foot = section.footer.paragraphs[0]
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_after = Pt(0)
    left = foot.add_run("%s  ·  Trang " % meta.get("classification", ""))
    left.font.size = Pt(8)
    left.font.color.rgb = rgb("808080")
    add_field(foot, "PAGE", "1").font.size = Pt(8)
    mid = foot.add_run(" / ")
    mid.font.size = Pt(8)
    mid.font.color.rgb = rgb("808080")
    add_field(foot, "NUMPAGES", "1").font.size = Pt(8)
    for r in foot.runs:
        r.font.name = BODY_FONT
        r.font.color.rgb = rgb("808080")
        r.font.size = Pt(8)


# ---------------------------------------------------------------- tables
def force_fixed_layout(table, widths_in, total_in):
    """Trap 4: autofit lets one long cell squeeze every other column to nothing.

    Fixed layout needs three things to agree: w:tblLayout, w:tblGrid, and the width on every
    single cell. Setting only the column object is not enough - python-docx writes it per cell.
    """
    total = float(sum(widths_in)) or 1.0
    scale = total_in / total
    widths = [max(0.35, w * scale) for w in widths_in]

    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ("w:tblLayout", "w:tblW"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tblPr.append(tblW)

    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(col)
    tbl.insert(1, grid)

    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])
    return widths


def repeat_header_row(table):
    trPr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


def write_cell(cell, text, *, bold=False, white=False, size=9.5):
    """Trap 6: cell.text = "" leaves an empty run that inherits Normal.

    So the value always goes into a run we created ourselves, and the paragraph formatting is
    reset for the table context (no justification, no space-after).
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.line_spacing = 1.1
    lines = str(text).split("\n")
    run = para.add_run(lines[0])
    for extra in lines[1:]:
        run.add_break()
        run.add_text(extra)
    run.font.size = Pt(size)
    run.font.name = BODY_FONT
    run.font.bold = bold
    if white:
        run.font.color.rgb = rgb("FFFFFF")


# ---------------------------------------------------------------- blocks
FLOW_MARK = [
    ("if ", "<>"),         # decision - a red diamond once D4 draws it
    ("return ", "()"),     # terminator - a green pill
    ("db:", "[/]"),        # data store - a parallelogram
]


def flow_marker(step):
    low = step.strip().lower()
    for prefix, mark in FLOW_MARK:
        if low.startswith(prefix):
            return mark
    return "[ ]"


class Renderer(object):
    def __init__(self, doc, counts):
        self.doc = doc
        self.counts = counts
        self.figure_no = 0        # continuous across the document (RULESET section 3)
        self.table_no = 0
        self.width_in = printable_width(doc)

    def paragraph(self, text, *, style=None, size=None, italic=False, color=None):
        p = self.doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.font.name = BODY_FONT
        if size:
            run.font.size = Pt(size)
        run.font.italic = italic
        if color:
            run.font.color.rgb = rgb(color)
        return p

    def block(self, b):
        t = b.get("t")
        handler = getattr(self, "b_" + t, None)
        if handler is None:
            self.counts["unknownBlocks"] = self.counts.get("unknownBlocks", 0) + 1
            return
        handler(b)
        self.counts["blocks"] = self.counts.get("blocks", 0) + 1

    def b_p(self, b):
        self.paragraph(b.get("text", ""))

    def b_bullets(self, b):
        for item in b.get("items", []):
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.add_run(item).font.name = BODY_FONT

    def b_num(self, b):
        # Trap 5: "List Number" counts across the whole document, so steps that must restart at 1
        # inside every section are numbered by hand.
        for i, item in enumerate(b.get("items", []), 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lead = p.add_run("%d. " % i)
            lead.font.bold = True
            lead.font.name = BODY_FONT
            p.add_run(item).font.name = BODY_FONT

    def b_table(self, b):
        headers = b.get("headers") or []
        rows = b.get("rows") or []
        width = len(headers) or (len(rows[0]) if rows else 0)
        if not width:
            return
        self.table_no += 1
        table = self.doc.add_table(rows=1 + len(rows), cols=width)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, head in enumerate(headers or [""] * width):
            cell = table.rows[0].cells[i]
            write_cell(cell, head, bold=True, white=True)
            shade(cell._tc, ACCENT)
        for r, row in enumerate(rows):
            for c, value in enumerate(row[:width]):
                cell = table.rows[r + 1].cells[c]
                write_cell(cell, value)
                if r % 2 == 1:
                    shade(cell._tc, ZEBRA)

        widths = b.get("widths") or [1.0] * width
        if len(widths) != width:
            widths = [1.0] * width
        force_fixed_layout(table, widths, self.width_in)
        repeat_header_row(table)
        caption = b.get("caption")
        if caption:
            p = self.paragraph("Bảng %d. %s" % (self.table_no, caption), size=9, italic=True,
                               color="808080")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        self.counts["tables"] = self.counts.get("tables", 0) + 1

    def b_code(self, b):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.05
        shade(p._p, CODE_BG)
        lines = str(b.get("text", "")).split("\n")
        run = p.add_run(lines[0])
        for extra in lines[1:]:
            run.add_break()
            run.add_text(extra)
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        self.counts["code"] = self.counts.get("code", 0) + 1

    def figure_caption(self, caption, alt):
        self.figure_no += 1
        p = self.paragraph("Hình %d. %s" % (self.figure_no, caption or ""), size=9, italic=True,
                           color="808080")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if alt:
            # Until the picture itself exists, the alt text is the only description a reader (or a
            # screen reader) gets, so it is written into the document rather than dropped.
            q = self.paragraph("Mô tả hình: %s" % alt, size=8.5, italic=True, color="808080")
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def b_figure(self, b):
        src = b.get("src") or ""
        placed = False
        if src and os.path.isfile(src):
            try:
                self.doc.add_picture(src, width=Inches(min(6.2, self.width_in)))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                placed = True
            except Exception as exc:                                  # noqa: BLE001
                OUT["warnings"].append("Không chèn được hình %s: %s" % (src, exc))
        if not placed:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            shade(p._p, ZEBRA)
            run = p.add_run("[ chưa có hình — sơ đồ được vẽ ở bản sau ]")
            run.font.size = Pt(9)
            run.font.name = BODY_FONT
            run.font.color.rgb = rgb("808080")
        self.figure_caption(b.get("caption"), b.get("alt"))
        self.counts["figures"] = self.counts.get("figures", 0) + 1

    def b_flow(self, b):
        # D2 has no diagram engine yet, so a flow becomes the numbered list of its steps with the
        # shape convention spelled out. D4 replaces this with a drawing from flow.py.
        for i, step in enumerate(b.get("steps", []), 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lead = p.add_run("%d. %s " % (i, flow_marker(step)))
            lead.font.name = MONO_FONT
            lead.font.size = Pt(9)
            lead.font.color.rgb = rgb(ACCENT)
            p.add_run(step).font.name = BODY_FONT
        self.figure_caption(b.get("caption") or "Sơ đồ luồng (dạng liệt kê bước)", b.get("alt"))
        self.counts["flows"] = self.counts.get("flows", 0) + 1

    def b_refs(self, b):
        for pair in b.get("items", []):
            label = pair[0] if len(pair) > 0 else ""
            url = pair[1] if len(pair) > 1 else ""
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            if url:
                add_hyperlink(p, url, label or url)
            else:
                p.add_run(label).font.name = BODY_FONT

    def b_callout(self, b):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        colour = {"warn": "C07F17", "danger": "B02418"}.get(b.get("level"), ACCENT)
        shade(p._p, ZEBRA)
        left_border(p, colour)
        run = p.add_run(str(b.get("text", "")))
        run.font.name = BODY_FONT
        run.font.size = Pt(10.5)
        run.font.color.rgb = rgb("333333")


# ---------------------------------------------------------------- document assembly
def level_of(num):
    return min(4, max(1, str(num).count(".") + 1))


def front_matter(doc, meta):
    doc.add_paragraph(meta["title"], style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(meta.get("project") or "")
    run.font.size = Pt(12)
    run.font.color.rgb = rgb("555555")
    if meta.get("draft"):
        d = doc.add_paragraph()
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mark = d.add_run("BẢN NHÁP — nội dung chưa hoàn chỉnh")
        mark.font.size = Pt(11)
        mark.font.bold = True
        mark.font.color.rgb = rgb("C07F17")

    rows = [("Mã tài liệu", meta.get("docId") or "-"),
            ("Phiên bản", meta.get("version") or "-"),
            ("Phân loại", meta.get("classification") or "-"),
            ("Trạng thái", meta.get("docStatus") or "-"),
            ("Chuẩn áp dụng", meta.get("standard") or "-")]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        write_cell(table.rows[i].cells[0], k, bold=True)
        write_cell(table.rows[i].cells[1], v)
    force_fixed_layout(table, [1.3, 3.0], printable_width(doc))

    doc.add_paragraph()
    doc.add_paragraph("Mục lục", style="Heading 1")
    toc = doc.add_paragraph()
    add_field(toc, r'TOC \o "1-3" \h \z \u',
              "Mục lục sẽ hiện khi mở bằng Word (hoặc nhấn F9 để cập nhật).")
    doc.add_page_break()


def sources_line(renderer, ir):
    """N2 made visible: a reader can check every section against the code it came from."""
    files = []
    for src in (ir.get("sources") or []):
        f = src.get("file")
        if f and f not in files:
            files.append(f)
    for block in ir.get("blocks") or []:
        for src in (block.get("sources") or []):
            f = src.get("file")
            if f and f not in files:
                files.append(f)
    assumed = sum(1 for b in (ir.get("blocks") or []) if b.get("assumption"))
    owner = sum(1 for b in (ir.get("blocks") or []) if b.get("providedBy"))
    if not files and not assumed and not owner:
        return
    bits = []
    if files:
        bits.append("Nguồn: " + ", ".join(files[:8]) + (" …" if len(files) > 8 else ""))
    if assumed:
        bits.append("%d khối là suy luận chưa kiểm chứng" % assumed)
    if owner:
        bits.append("%d khối do chủ sản phẩm cung cấp" % owner)
    p = renderer.paragraph("  ·  ".join(bits), size=8.5, italic=True, color="808080")
    p.paragraph_format.space_before = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def render_document(spec, dest_dir, draft):
    meta = {
        "title": spec.get("title") or "Tài liệu",
        "project": spec.get("project") or "",
        "docId": spec.get("docId") or "",
        "version": spec.get("version") or "",
        "classification": spec.get("classification") or "",
        "docStatus": spec.get("docStatus") or "",
        "standard": spec.get("standard") or "",
        "draft": bool(draft),
    }
    # normpath so the reported path is the one the user can paste into Explorer, not a mix of
    # separators from joining a POSIX-ish destDir with a Windows filename.
    path = os.path.normpath(os.path.join(dest_dir, spec.get("file") or (meta["title"] + ".docx")))

    # Trap 3: a file open in Word cannot be written. Find out before spending the work, and report
    # the name instead of raising - the other documents of the same export must still go out.
    if os.path.exists(path):
        try:
            with open(path, "r+b"):
                pass
        except (PermissionError, OSError) as exc:                    # noqa: BLE001
            return {"skipped": {"path": path, "reason":
                    "Tệp đang mở trong Word hoặc bị khoá (%s). Đóng tệp rồi xuất lại." % exc.__class__.__name__}}

    doc = Document()
    apply_theme(doc)
    build_header_footer(doc, meta)
    front_matter(doc, meta)

    counts = {"sections": 0, "empty": 0}
    renderer = Renderer(doc, counts)
    for section in spec.get("sections") or []:
        num = section.get("num") or ""
        heading = ("%s. %s" % (num, section.get("title") or "")).strip()
        doc.add_paragraph(heading, style="Heading %d" % level_of(num))
        ir = section.get("ir")
        if not ir or not (ir.get("blocks") or []):
            renderer.paragraph("Mục này chưa có nội dung.", size=10.5, italic=True, color="808080")
            counts["empty"] += 1
            continue
        for block in ir["blocks"]:
            renderer.block(block)
        sources_line(renderer, ir)
        counts["sections"] += 1

    try:
        doc.save(path)
    except (PermissionError, OSError) as exc:                        # noqa: BLE001
        return {"skipped": {"path": path, "reason":
                "Không ghi được tệp (%s): %s" % (exc.__class__.__name__, exc)}}
    counts["figures"] = renderer.figure_no
    counts["tableCaptions"] = renderer.table_no
    return {"file": {"path": path, "bytes": os.path.getsize(path), "counts": counts,
                     "title": meta["title"]}}


def main():
    args = sys.argv[1:]
    if "--check" in args:
        import docx
        OUT["ok"] = True
        OUT["python"] = sys.version.split()[0]
        OUT["pythonDocx"] = getattr(docx, "__version__", "unknown")
        sys.stdout.write(json.dumps(OUT, ensure_ascii=False))
        return
    if not args:
        die("Thiếu đường dẫn tệp payload JSON.")
    try:
        with open(args[0], "r", encoding="utf-8") as fh:
            payload = json.load(fh)
    except Exception as exc:                                          # noqa: BLE001
        die("Không đọc được payload: %s" % exc)

    dest = payload.get("destDir") or os.getcwd()
    try:
        os.makedirs(dest, exist_ok=True)
    except OSError as exc:
        die("Không tạo được thư mục đích %s: %s" % (dest, exc), kind="dest")

    draft = bool(payload.get("draft"))
    total = {}
    for spec in payload.get("docs") or []:
        result = render_document(spec, dest, draft)
        if "skipped" in result:
            OUT["skipped"].append(result["skipped"])
            continue
        OUT["files"].append(result["file"])
        for k, v in result["file"]["counts"].items():
            if isinstance(v, int):
                total[k] = total.get(k, 0) + v
    OUT["counts"] = total
    OUT["ok"] = bool(OUT["files"]) or not (payload.get("docs") or [])
    if not OUT["files"] and OUT["skipped"]:
        OUT["error"] = "Không xuất được tệp nào: tất cả đang bị mở hoặc bị khoá."
    sys.stdout.write(json.dumps(OUT, ensure_ascii=False))


if __name__ == "__main__":
    # Any unexpected failure still leaves ONE json object on stdout: the Node side parses stdout
    # and a raw traceback there would surface as "không đọc được kết quả" instead of the reason.
    try:
        main()
    except SystemExit:
        raise
    except Exception as exc:                                          # noqa: BLE001
        import traceback
        die("Lỗi không lường trước khi dựng .docx: %s\n%s"
            % (exc, traceback.format_exc(limit=4)))
